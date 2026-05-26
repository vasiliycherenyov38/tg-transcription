import os
import logging
import tempfile
import subprocess
import httpx
import re
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, MessageHandler, CallbackQueryHandler, filters, ContextTypes
from http.server import HTTPServer, BaseHTTPRequestHandler
import threading
from docx import Document as DocxDocument
from docx.shared import Pt
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

TELEGRAM_TOKEN = os.environ["TELEGRAM_TOKEN"]
GROQ_API_KEY = os.environ["GROQ_API_KEY"]

SUPPORTED_EXTENSIONS = [".mp3", ".mp4", ".wav", ".m4a", ".ogg", ".webm", ".mpeg", ".mpga"]

# Временное хранилище транскриптов
transcripts = {}

def convert_to_audio(file_path: str) -> str:
    audio_path = file_path + ".mp3"
    subprocess.run([
        "ffmpeg", "-i", file_path,
        "-vn", "-acodec", "mp3",
        "-ab", "64k", "-ar", "16000",
        "-y", audio_path
    ], capture_output=True)
    return audio_path

def split_audio(file_path: str, chunk_minutes: int = 10) -> list:
    chunk_seconds = chunk_minutes * 60
    duration_result = subprocess.run([
        "ffprobe", "-v", "quiet",
        "-show_entries", "format=duration",
        "-of", "csv=p=0", file_path
    ], capture_output=True, text=True)
    
    try:
        duration = float(duration_result.stdout.strip())
    except:
        return [(file_path, 0)]
    
    if duration <= chunk_seconds:
        return [(file_path, 0)]
    
    chunks = []
    start = 0
    index = 0
    while start < duration:
        chunk_path = file_path + f"_chunk{index}.mp3"
        subprocess.run([
            "ffmpeg", "-i", file_path,
            "-ss", str(start),
            "-t", str(chunk_seconds),
            "-acodec", "mp3", "-ab", "64k", "-ar", "16000",
            "-y", chunk_path
        ], capture_output=True)
        chunks.append((chunk_path, start))
        start += chunk_seconds
        index += 1
    
    return chunks

async def transcribe_chunk(file_path: str, offset_seconds: float = 0) -> str:
    url = "https://api.groq.com/openai/v1/audio/transcriptions"
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    
    with open(file_path, "rb") as f:
        files = {
            "file": (os.path.basename(file_path), f, "audio/mpeg"),
            "model": (None, "whisper-large-v3"),
            "language": (None, "ru"),
            "response_format": (None, "verbose_json"),
        }
        async with httpx.AsyncClient(timeout=300) as client:
            response = await client.post(url, headers=headers, files=files)
    
    if response.status_code != 200:
        raise Exception(f"Groq error: {response.text}")
    
    data = response.json()
    segments = data.get("segments", [])
    
    if not segments:
        return data.get("text", "")
    
    lines = []
    for seg in segments:
        start = int(seg["start"] + offset_seconds)
        minutes = start // 60
        seconds = start % 60
        timestamp = f"[{minutes:02d}:{seconds:02d}]"
        lines.append(f"{timestamp} {seg['text'].strip()}")
    
    return "\n".join(lines)

async def process_file(file_path: str, message) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext in [".mp4", ".webm", ".mpeg", ".mov"]:
        await message.reply_text("🎬 Извлекаю аудио из видео...")
        audio_path = convert_to_audio(file_path)
        os.unlink(file_path)
    else:
        audio_path = file_path
    
    chunks = split_audio(audio_path)
    
    if len(chunks) > 1:
        total = len(chunks)
        await message.reply_text(f"✂️ Нарезал на {total} частей по 10 минут, транскрибирую...")
        
        all_text = []
        for i, (chunk_path, offset) in enumerate(chunks):
            await message.reply_text(f"⏳ Часть {i+1}/{total}...")
            text = await transcribe_chunk(chunk_path, offset)
            all_text.append(text)
            os.unlink(chunk_path)
        
        if audio_path != file_path:
            os.unlink(audio_path)
        return "\n".join(all_text)
    else:
        chunk_path, offset = chunks[0]
        text = await transcribe_chunk(chunk_path, offset)
        os.unlink(audio_path)
        return text

async def download_from_drive(url: str) -> tuple:
    async with httpx.AsyncClient(timeout=600, follow_redirects=True) as client:
        r = await client.get(url)
        
        if b"confirm=" in r.content:
            match = re.search(rb"confirm=([0-9A-Za-z_]+)", r.content)
            if match:
                confirm = match.group(1).decode()
                file_id = url.split("id=")[1] if "id=" in url else ""
                url = f"https://drive.google.com/uc?export=download&confirm={confirm}&id={file_id}"
                r = await client.get(url)
        
        content_type = r.headers.get("content-type", "")
        ext = ".wav"
        if "mp4" in content_type or "video" in content_type:
            ext = ".mp4"
        elif "ogg" in content_type:
            ext = ".ogg"
        elif "mp3" in content_type or "mpeg" in content_type:
            ext = ".mp3"
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(r.content)
            return tmp.name, ext

def make_txt(text: str, filename: str) -> str:
    path = tempfile.mktemp(suffix=".txt")
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"Транскрипт: {filename}\n")
        f.write(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
        f.write("=" * 50 + "\n\n")
        f.write(text)
    return path

def make_docx(text: str, filename: str) -> str:
    doc = DocxDocument()
    
    title = doc.add_heading("Транскрипт", 0)
    title.alignment = 1
    
    doc.add_paragraph(f"Файл: {filename}")
    doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    doc.add_paragraph("")
    
    for line in text.split("\n"):
        if line.strip():
            p = doc.add_paragraph()
            if line.startswith("["):
                # Строка с таймкодом
                bracket_end = line.find("]")
                if bracket_end > 0:
                    timestamp = line[:bracket_end+1]
                    content = line[bracket_end+1:].strip()
                    run_ts = p.add_run(timestamp + " ")
                    run_ts.bold = True
                    run_ts.font.size = Pt(9)
                    run_content = p.add_run(content)
                    run_content.font.size = Pt(11)
            else:
                run = p.add_run(line)
                run.font.size = Pt(11)
    
    path = tempfile.mktemp(suffix=".docx")
    doc.save(path)
    return path

async def send_format_choice(message, transcript_text: str, source_name: str):
    chat_id = message.chat_id
    transcripts[chat_id] = {
        "text": transcript_text,
        "name": source_name
    }
    
    keyboard = [
        [
            InlineKeyboardButton("📄 TXT", callback_data="format_txt"),
            InlineKeyboardButton("📝 DOCX", callback_data="format_docx"),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await message.reply_text(
        "✅ Транскрипция готова! В каком формате отправить?",
        reply_markup=reply_markup
    )

async def handle_format_choice(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    
    chat_id = query.message.chat_id
    data = transcripts.get(chat_id)
    
    if not data:
        await query.edit_message_text("Транскрипт не найден, отправь файл заново.")
        return
    
    text = data["text"]
    name = data["name"]
    fmt = query.data
    
    await query.edit_message_text("⏳ Формирую файл...")
    
    if fmt == "format_txt":
        file_path = make_txt(text, name)
        with open(file_path, "rb") as f:
            await query.message.reply_document(f, filename=f"transcript_{datetime.now().strftime('%d%m%Y_%H%M')}.txt")
        os.unlink(file_path)
    
    elif fmt == "format_docx":
        file_path = make_docx(text, name)
        with open(file_path, "rb") as f:
            await query.message.reply_document(f, filename=f"transcript_{datetime.now().strftime('%d%m%Y_%H%M')}.docx")
        os.unlink(file_path)
    
    del transcripts[chat_id]

async def handle_audio(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = update.message
    await message.reply_text("⏳ Получил файл, начинаю обработку...")
    
    try:
        audio = message.audio or message.voice or message.video or message.document
        
        if message.document:
            fname = message.document.file_name or ""
            if not any(fname.lower().endswith(ext) for ext in SUPPORTED_EXTENSIONS):
                await message.reply_text("Формат не поддерживается. Отправь mp3, mp4, wav, m4a, ogg.")
                return
            source_name = fname
        else:
            source_name = f"audio_{datetime.now().strftime('%d%m%Y_%H%M')}"
        
        tg_file = await audio.get_file()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".ogg") as tmp:
            tmp_path = tmp.name
        
        await tg_file.download_to_drive(tmp_path)
        result = await process_file(tmp_path, message)
        await send_format_choice(message, result, source_name)
    
    except Exception as e:
        logger.error(f"Error: {e}")
        await message.reply_text(f"Ошибка: {str(e)}")

async def handle_text(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = update.message
    text = message.text.strip()
    
    if not text.startswith("http"):
        await message.reply_text(
            "Привет! Отправь мне:\n"
            "• Аудио или голосовое сообщение\n"
            "• Видео или файл\n"
            "• Ссылку на Google Drive"
        )
        return
    
    await message.reply_text("⏳ Скачиваю файл...")
    
    try:
        if "drive.google.com" in text:
            if "/file/d/" in text:
                file_id = text.split("/file/d/")[1].split("/")[0]
                text = f"https://drive.google.com/uc?export=download&id={file_id}"
            tmp_path, _ = await download_from_drive(text)
        else:
            async with httpx.AsyncClient(timeout=600, follow_redirects=True) as client:
                response = await client.get(text)
            
            content_type = response.headers.get("content-type", "")
            ext = ".mp3"
            if "mp4" in content_type or "video" in content_type:
                ext = ".mp4"
            elif "ogg" in content_type:
                ext = ".ogg"
            elif "wav" in content_type:
                ext = ".wav"
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                tmp.write(response.content)
                tmp_path = tmp.name
        
        await message.reply_text("✅ Файл скачан, обрабатываю...")
        result = await process_file(tmp_path, message)
        await send_format_choice(message, result, "google_drive")
    
    except Exception as e:
        logger.error(f"Error: {e}")
        await message.reply_text(f"Ошибка: {str(e)}")

class HealthHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        self.send_response(200)
        self.end_headers()
        self.wfile.write(b"OK")
    def log_message(self, format, *args):
        pass

def run_health_server():
    server = HTTPServer(("0.0.0.0", 8080), HealthHandler)
    server.serve_forever()

async def post_init(application):
    await application.bot.delete_webhook(drop_pending_updates=True)

def main():
    threading.Thread(target=run_health_server, daemon=True).start()
    app = Application.builder().token(TELEGRAM_TOKEN).post_init(post_init).build()
    app.add_handler(MessageHandler(filters.AUDIO | filters.VOICE | filters.VIDEO | filters.Document.ALL, handle_audio))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(CallbackQueryHandler(handle_format_choice))
    app.run_polling(drop_pending_updates=True)

if __name__ == "__main__":
    main()
